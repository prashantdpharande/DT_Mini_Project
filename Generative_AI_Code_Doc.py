import os
import requests
import openai
import ast
from docx import Document
from flask import Flask, request, jsonify

# 🔹 Replace with your credentials
GITHUB_TOKEN = "your_github_token"
OPENAI_API_KEY = "your_openai_api_key"
REPO_OWNER = "repo_owner"
REPO_NAME = "repo_name"

# 🔹 OpenAI API setup
openai.api_key = OPENAI_API_KEY

def fetch_repo_contents():
    """Fetches all Python files from the GitHub repository."""
    url = f"https://api.github.com/repos/{REPO_OWNER}/{REPO_NAME}/git/trees/main?recursive=1"
    headers = {"Authorization": f"token {GITHUB_TOKEN}"}
    response = requests.get(url, headers=headers)
    
    if response.status_code == 200:
        repo_data = response.json()
        return [file["path"] for file in repo_data["tree"] if file["path"].endswith(".py")]
    else:
        print(f"Failed to fetch repository: {response.json()}")
        return []

def fetch_file_content(file_path):
    """Fetches raw content of a Python file from GitHub."""
    url = f"https://raw.githubusercontent.com/{REPO_OWNER}/{REPO_NAME}/main/{file_path}"
    response = requests.get(url)
    return response.text if response.status_code == 200 else ""

def extract_code_structure(code):
    """Parses Python code to extract functions, classes, and docstrings."""
    tree = ast.parse(code)
    extracted_data = {"classes": [], "functions": []}

    for node in ast.walk(tree):
        if isinstance(node, ast.ClassDef):
            docstring = ast.get_docstring(node) or "No docstring provided."
            extracted_data["classes"].append((node.name, docstring))
        elif isinstance(node, ast.FunctionDef):
            docstring = ast.get_docstring(node) or "No docstring provided."
            extracted_data["functions"].append((node.name, docstring))

    return extracted_data

def generate_ai_description(code, file_path):
    """Uses OpenAI GPT to generate a meaningful description of the file."""
    prompt = f"""
    You are an expert code analyst. Analyze the following Python code and provide a structured summary.

    - Explain its purpose
    - List key classes and functions with their roles
    - Mention any dependencies or important logic
    
    Code from {file_path}:
    ```python
    {code}
    ```
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "system", "content": "You are an AI code documentation assistant."},
                  {"role": "user", "content": prompt}]
    )
    return response["choices"][0]["message"]["content"]
def structure_code_1(code):
    """Parses Python code to extract functions, classes, and docstrings."""
    tree = ast.parse(code)
    extracted_data = {"classes": [], "functions": []}

    for node in ast.walk(tree):
        if isinstance(node, ast.ClassDef):
            docstring = ast.get_docstring(node) or "No docstring provided."
            extracted_data["classes"].append((node.name, docstring))
        elif isinstance(node, ast.FunctionDef):
            docstring = ast.get_docstring(node) or "No docstring provided."
            extracted_data["functions"].append((node.name, docstring))

    return extracted_data
def generate_docx():
    """Generates a structured Word document with AI-generated documentation."""
    doc = Document()
    doc.add_heading(f'Documentation for {REPO_NAME}', 0)
    
    py_files = fetch_repo_contents()
    
    for file_path in py_files:
        code = fetch_file_content(file_path)
        if not code:
            continue
        
        structure = extract_code_structure(code)
        ai_description = generate_ai_description(code, file_path)
        
        doc.add_heading(file_path, level=1)
        doc.add_paragraph(ai_description, style="Normal")

        if structure["classes"]:
            doc.add_heading("Classes", level=2)
            for cls_name, docstring in structure["classes"]:
                doc.add_paragraph(f"🔹 {cls_name}: {docstring}")

        if structure["functions"]:
            doc.add_heading("Functions", level=2)
            for func_name, docstring in structure["functions"]:
                doc.add_paragraph(f"🔹 {func_name}: {docstring}")

    doc.save(f"{REPO_NAME}_documentation.docx")
    print(f"✅ Documentation generated: {REPO_NAME}_documentation.docx")

# 🔹 Flask Webhook for Real-Time Updates
app = Flask(__name__)

@app.route('/github-webhook', methods=['POST'])
def webhook():
    """Triggers documentation generation when a push event occurs."""
    data = request.json
    if "pusher" in data:
        print("🔄 Change detected in repo. Regenerating documentation...")
        generate_docx()
        return jsonify({"message": "Documentation updated"}), 200
    return jsonify({"message": "Ignored event"}), 400

if __name__ == "__main__":
    generate_docx()  # Run once to generate documentation initially
    app.run(port=5000)
